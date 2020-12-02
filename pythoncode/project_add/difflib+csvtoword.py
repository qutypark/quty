
# 목적: csv문서를 서로 비교 -> 비교결과 및 플래그를 word로 export

#=============================== 0.import library===============================

import pandas as pd
import numpy as np
import os, collections, csv


import difflib
import glob                        
import re

# word
from docx import Document        
from docx.shared import RGBColor #Helps to specify font Color
from docx.enum.text import WD_COLOR_INDEX


#=============================== １. read csvfile===============================

# 비교대상 데이터
csv_path_glob = os.path.join("localpath/*.csv") 
csv_paths = sorted(glob.glob(csv_path_glob)) 

# 비교기준 데이터
stdpath="standard.csv"

#=============================== 2.  비교 룰 함수===============================

def inline_diff(a, b):
    matcher = difflib.SequenceMatcher(None, a, b)
    def process_tag(tag, i1, i2, j1, j2):
        if tag == 'replace':
            return '{'  + matcher.a[i1:i2] + ' -> ' + matcher.b[j1:j2] + '}'
        if tag == 'delete':
            return '{- ' + matcher.a[i1:i2] + '}'
        if tag == 'equal':
            return matcher.a[i1:i2]
        if tag == 'insert':
            return '{+ ' + matcher.b[j1:j2] + '}'
        assert False, "Unknown tag %r"%tag
    return ''.join(process_tag(*t) for t in matcher.get_opcodes())


#=============================3. csv라인별로 비교 함수적용=====================
result_list0=[None]*len(csv_paths)
lineh=[]
linet0=[None]*len(csv_paths)

with open(stdpath,"r",encoding=cp932) as h:
    for line in h:
        line=re.findall('.*?[。]', line)
        lineh.extend(line)
for i,cpath in enumerate(csv_paths):
    linet=[]
    result_list=[]
    with open(cpath, "r",encoding=cp932) as my_input_file:
        for line in my_input_file:
            line=re.findall('.*?[。]', line)
            linet.extend(line)
        linet0[i]=linet
    for itemh,itemt in zip(lineh,linet0[i]):
        result=inline_diff(itemh,itemt)
        result_list.append("".join(result))
    result_list0[i]=result_list


# =======4.  비교결과를 word문서로 저장/ 색깔 및 조건 별 플래그 생성=======

from operator import itemgetter
for i, cpath in enumerate(csv_paths):
    mydoc=Document()
    #저장 시의 word파일 명 생성=원래 csv파일 명
    x=re.split("[/.]",cpath)[-2]
    y=x.replace("\\",".")
    z=re.split("[..]",y)[-1]
    resultname = '%s.docx' % (z)
    result_path="C:/local/result/"
    result_file= os.path.join(result_path,resultname) 
    if "{" in itemgetter(2)(result_list0[i]):  # example of flag generation:  "{" 가 특정위치에 있다면 flag1생성
        run = mydoc.add_heading().add_run("flag1") # 플래그를 워드 문서 헤드화
        run.font.color.rgb = RGBColor(0, 0, 255)  # 플래그 색상 blue
    else:
        run = mydoc.add_heading().add_run("flag2")   # example of flag generation:"{" 가 특정위치에 없다면 flag2생성
        run.font.color.rgb = RGBColor(0, 0, 0)   # 플래그 색상 black      
    for item in result_list0[i]:
        p=mydoc.add_paragraph()
        if "{" in item:
            for j,item1 in enumerate(item):
                a=re.findall(r"{",item1)
                b=re.findall(r"}",item1)
                run = p.add_run(item1)
                if(a):run.font.highlight_color = WD_COLOR_INDEX.YELLOW   # 비교 차이 부분: 하이라이트 색상지정  
                elif(b):run.font.highlight_color = WD_COLOR_INDEX.YELLOW  # 비교 차이 부분: 하이라이트 색상지정  
                else:run.font.color.rgb = RGBColor(0, 0, 0)  # 비교 차이 부분이 없다면 글자 색 black 
        else:
            run = p.add_run(item)
            run.font.color.rgb = RGBColor(0, 0, 0)
    mydoc.save(result_file)             


